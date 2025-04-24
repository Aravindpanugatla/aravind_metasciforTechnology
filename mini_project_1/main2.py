import streamlit as st
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from collections import defaultdict
from docx import Document

#Scraper Functions

visited = set()
site_map = defaultdict(list)

def clean_url(base, link):
    try:
        full_url = urljoin(base, link)
        if urlparse(full_url).netloc == urlparse(base).netloc:
            return full_url.split("#")[0]
    except:
        return None

def scrape_site(url):
    visited.clear()
    content_data = {}

    def crawl(current_url):
        if current_url in visited:
            return
        try:
            res = requests.get(current_url, timeout=8)
            if res.status_code != 200:
                return
            soup = BeautifulSoup(res.text, 'html.parser')
            visited.add(current_url)

            headings = [h.get_text(strip=True) for h in soup.find_all(['h1', 'h2', 'h3'])]
            paragraphs = [p.get_text(strip=True) for p in soup.find_all('p')]
            links = [clean_url(current_url, a['href']) for a in soup.find_all('a', href=True)]
            links = list(filter(None, links))

            site_map[current_url] = links
            content_data[current_url] = {
                "headings": headings[:5],
                "paragraphs": paragraphs[:5],
                "links": links
            }

            for link in links:
                crawl(link)

        except Exception as e:
            print(f"Error while scraping {current_url}: {e}")
            pass

    crawl(url)
    return content_data

def generate_report(content_data, filename="webscraper_Report.docx"):
    doc = Document()
    doc.add_heading(' Webscraping Report', 0)
    for url, data in content_data.items():
        doc.add_heading(f"URL: {url}", level=1)

        doc.add_heading("Headings", level=2)
        for h in data['headings']:
            doc.add_paragraph(h)

        doc.add_heading("Paragraphs", level=2)
        for p in data['paragraphs']:
            doc.add_paragraph(p)

        doc.add_heading("Internal Links", level=2)
        for l in data['links']:
            doc.add_paragraph(l)

    doc.save(filename)
    return filename

#Streamlit UI

st.set_page_config(page_title="SiteSketcher", layout="wide")
st.title(" Website Scraper")
st.markdown("Scrape any website (including subpages) and export a structured Word document report.")

url = st.text_input("🔗 Enter Website URL", "https://example.com")

if st.button(" Scrape and Generate Report"):
    with st.spinner("Scraping site... please wait"):
        data = scrape_site(url)
        if data:
            filename = generate_report(data)
            with open(filename, "rb") as f:
                st.success("Report generated successfully!")
                st.download_button(
                    label="Download Report (.docx)",
                    data=f,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning(" No content could be extracted from this site.")
